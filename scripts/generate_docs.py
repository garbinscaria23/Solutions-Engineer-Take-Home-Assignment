# scripts/generate_docs.py
import os
import json

# docx imports
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# reportlab imports
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PDF_PATH = os.path.join(PROJECT_ROOT, "api_documentation.pdf")
DOCX_PATH = os.path.join(PROJECT_ROOT, "api_documentation.docx")

# API Doc data
endpoints = [
    {
        "title": "1. Ingest Events",
        "method": "POST",
        "path": "/events",
        "desc": "Accepts payment lifecycle events (payment_initiated, payment_processed, payment_failed, settled). Event ingestion is idempotent, duplicate events (same event_id) return success without mutating transaction history.",
        "params": [],
        "req_body": {
            "event_id": "b768e3a7-9eb3-4603-b21c-a54cc95661bc",
            "event_type": "payment_initiated",
            "transaction_id": "2f86e94c-239c-4302-9874-75f28e3474ee",
            "merchant_id": "merchant_2",
            "merchant_name": "FreshBasket",
            "amount": 15248.29,
            "currency": "INR",
            "timestamp": "2026-01-08T12:11:58.085567+00:00"
        },
        "resp_body": {
            "event_id": "b768e3a7-9eb3-4603-b21c-a54cc95661bc",
            "event_type": "payment_initiated",
            "amount": "15248.29",
            "currency": "INR",
            "timestamp": "2026-01-08T12:11:58.085567+00:00",
            "received_at": "2026-07-19T17:37:46.429153+05:30"
        },
        "curl": "curl -X POST http://localhost:8000/events \\\n  -H \"Content-Type: application/json\" \\\n  -d '{\n    \"event_id\": \"b768e3a7-9eb3-4603-b21c-a54cc95661bc\",\n    \"event_type\": \"payment_initiated\",\n    \"transaction_id\": \"2f86e94c-239c-4302-9874-75f28e3474ee\",\n    \"merchant_id\": \"merchant_2\",\n    \"merchant_name\": \"FreshBasket\",\n    \"amount\": 15248.29,\n    \"currency\": \"INR\",\n    \"timestamp\": \"2026-01-08T12:11:58.085567+00:00\"\n  }'"
    },
    {
        "title": "2. List Transactions",
        "method": "GET",
        "path": "/transactions",
        "desc": "Retrieves transactions with support for advanced filtering, sorting, pagination, and date-range conditions.",
        "params": [
            ("merchant_id", "string", "No", "Filter by Merchant ID"),
            ("status", "string", "No", "Filter by status (payment_initiated, payment_processed, payment_failed, settled)"),
            ("start_date", "datetime", "No", "Filter created on or after ISO timestamp"),
            ("end_date", "datetime", "No", "Filter created on or before ISO timestamp"),
            ("skip", "integer", "No", "Offset for pagination (default: 0)"),
            ("limit", "integer", "No", "Limit for pagination (default: 100)"),
            ("sort_by", "string", "No", "Sort by field: created_at, updated_at, amount, status (default: created_at)"),
            ("sort_order", "string", "No", "Sort direction: asc, desc (default: desc)")
        ],
        "req_body": None,
        "resp_body": [
            {
                "id": "2f86e94c-239c-4302-9874-75f28e3474ee",
                "merchant_id": "merchant_2",
                "merchant_name": "FreshBasket",
                "amount": "15248.29",
                "currency": "INR",
                "status": "payment_failed",
                "has_discrepancy": False,
                "discrepancy_reason": None,
                "created_at": "2026-01-08T12:11:58.085567+05:30",
                "updated_at": "2026-01-08T12:38:58.085567+05:30"
            }
        ],
        "curl": "curl -X GET \"http://localhost:8000/transactions?merchant_id=merchant_2&status=payment_failed&limit=2\""
    },
    {
        "title": "3. Fetch Transaction Details",
        "method": "GET",
        "path": "/transactions/{transaction_id}",
        "desc": "Returns complete details for a single transaction, including its associated merchant data and its full chronological event history/state transitions.",
        "params": [
            ("transaction_id", "string", "Yes", "Path Parameter. The unique transaction identifier.")
        ],
        "req_body": None,
        "resp_body": {
            "id": "2f86e94c-239c-4302-9874-75f28e3474ee",
            "merchant_id": "merchant_2",
            "merchant_name": "FreshBasket",
            "amount": "15248.29",
            "currency": "INR",
            "status": "payment_failed",
            "has_discrepancy": False,
            "discrepancy_reason": None,
            "created_at": "2026-01-08T12:11:58.085567+05:30",
            "updated_at": "2026-01-08T12:38:58.085567+05:30",
            "events": [
                {
                    "event_id": "b768e3a7-9eb3-4603-b21c-a54cc95661bc",
                    "event_type": "payment_initiated",
                    "amount": "15248.29",
                    "currency": "INR",
                    "timestamp": "2026-01-08T12:11:58.085567+05:30",
                    "received_at": "2026-07-19T17:37:46.429153+05:30"
                },
                {
                    "event_id": "da46895f-4b47-4505-900e-d067f64a55eb",
                    "event_type": "payment_failed",
                    "amount": "15248.29",
                    "currency": "INR",
                    "timestamp": "2026-01-08T12:38:58.085567+05:30",
                    "received_at": "2026-07-19T17:37:46.574057+05:30"
                }
            ]
        },
        "curl": "curl -X GET http://localhost:8000/transactions/2f86e94c-239c-4302-9874-75f28e3474ee"
    },
    {
        "title": "4. Reconciliation Summary",
        "method": "GET",
        "path": "/reconciliation/summary",
        "desc": "Returns aggregated transaction volume, value, and discrepancy counts grouped by dimensions: merchant, calendar date, and state status.",
        "params": [],
        "req_body": None,
        "resp_body": {
            "by_merchant": [
                {
                    "dimension": "merchant",
                    "group_value": "FreshBasket (merchant_2)",
                    "total_transactions": 810,
                    "total_amount": "21040319.45",
                    "discrepancy_count": 52
                }
            ],
            "by_date": [
                {
                    "dimension": "date",
                    "group_value": "2026-01-08",
                    "total_transactions": 250,
                    "total_amount": "5204810.15",
                    "discrepancy_count": 12
                }
            ],
            "by_status": [
                {
                    "dimension": "status",
                    "group_value": "settled",
                    "total_transactions": 2341,
                    "total_amount": "61023049.20",
                    "discrepancy_count": 48
                }
            ]
        },
        "curl": "curl -X GET http://localhost:8000/reconciliation/summary"
    },
    {
        "title": "5. Reconciliation Discrepancies",
        "method": "GET",
        "path": "/reconciliation/discrepancies",
        "desc": "Returns transactions flagged with status/amount inconsistencies. Discrepancies include: settlements on failed payments, processed events but never settled (>6 hrs), duplicate conflicting events, and amount/currency mismatches.",
        "params": [
            ("skip", "integer", "No", "Offset for pagination (default: 0)"),
            ("limit", "integer", "No", "Limit for pagination (default: 100)")
        ],
        "req_body": None,
        "resp_body": [
            {
                "id": "02d878de-f807-4bf1-9b16-7098be6e54fe",
                "merchant_id": "merchant_5",
                "merchant_name": "StyleHub",
                "amount": "41263.87",
                "currency": "INR",
                "status": "settled",
                "has_discrepancy": True,
                "discrepancy_reason": "Duplicate conflicting events of type 'settled'",
                "created_at": "2026-01-08T14:29:10.700477+05:30",
                "updated_at": "2026-01-08T22:10:10.700477+05:30",
                "events": []
            }
        ],
        "curl": "curl -X GET \"http://localhost:8000/reconciliation/discrepancies?limit=2\""
    }
]

def generate_docx():
    print(f"Generating DOCX at {DOCX_PATH}...")
    doc = Document()
    
    # Base Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10.5)
    
    # Title Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Setu Reconciliation API Documentation\n")
    title_run.font.size = Pt(24)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(43, 85, 144) # Slate Blue
    
    subtitle_run = title_p.add_run("Developer API Reference & Integration Guide\n")
    subtitle_run.font.size = Pt(14)
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = RGBColor(120, 120, 120)
    
    doc.add_paragraph("\n" * 3)
    
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.add_run("Backend Service: FastAPI\nDatabase: PostgreSQL\nDate: July 2026\nVersion: 1.0.0\n")
    
    doc.add_page_break()
    
    # Intro
    h_intro = doc.add_heading("Overview", level=1)
    h_intro.runs[0].font.color.rgb = RGBColor(43, 85, 144)
    doc.add_paragraph(
        "This service provides payment lifecycle event ingestion and reconciliation capabilities. "
        "It supports event logging with strict idempotency check, transaction state maintenance, "
        "and discrepancy checks for reporting. This document describes the schema structure, request "
        "payloads, query parameters, responses, and curl commands for all active API endpoints."
    )
    
    # Base URL
    doc.add_heading("Base URL", level=2)
    doc.add_paragraph("Local Server: http://localhost:8000")
    
    doc.add_page_break()
    
    # API endpoints
    for ep in endpoints:
        h = doc.add_heading(ep["title"], level=1)
        h.runs[0].font.color.rgb = RGBColor(43, 85, 144)
        
        # Method and Path
        p_path = doc.add_paragraph()
        r_method = p_path.add_run(f"[{ep['method']}] ")
        r_method.bold = True
        r_method.font.color.rgb = RGBColor(200, 50, 50) if ep['method'] == "POST" else RGBColor(50, 150, 50)
        
        r_path = p_path.add_run(ep['path'])
        r_path.bold = True
        r_path.font.size = Pt(11.5)
        
        doc.add_paragraph(ep["desc"])
        
        # Query Parameters Table
        if ep["params"]:
            doc.add_heading("Query Parameters", level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Shading Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Parameter'
            hdr_cells[1].text = 'Type'
            hdr_cells[2].text = 'Required'
            hdr_cells[3].text = 'Description'
            for name, ptype, req, pdesc in ep["params"]:
                row_cells = table.add_row().cells
                row_cells[0].text = name
                row_cells[1].text = ptype
                row_cells[2].text = req
                row_cells[3].text = pdesc
            doc.add_paragraph() # Spacer
            
        # Request Body
        if ep["req_body"]:
            doc.add_heading("Request Payload (JSON)", level=2)
            code_block = doc.add_paragraph()
            code_block.paragraph_format.left_indent = Inches(0.25)
            r_code = code_block.add_run(json.dumps(ep["req_body"], indent=2))
            r_code.font.name = 'Consolas'
            r_code.font.size = Pt(9.5)
            
        # Response Body
        if ep["resp_body"]:
            doc.add_heading("Response Payload (JSON)", level=2)
            code_block = doc.add_paragraph()
            code_block.paragraph_format.left_indent = Inches(0.25)
            r_code = code_block.add_run(json.dumps(ep["resp_body"], indent=2))
            r_code.font.name = 'Consolas'
            r_code.font.size = Pt(9.5)
            
        # Curl Command
        doc.add_heading("Sample curl Command", level=2)
        curl_block = doc.add_paragraph()
        curl_block.paragraph_format.left_indent = Inches(0.25)
        r_curl = curl_block.add_run(ep["curl"])
        r_curl.font.name = 'Consolas'
        r_curl.font.size = Pt(9)
        r_curl.font.color.rgb = RGBColor(50, 50, 150)
        
        doc.add_page_break()
        
    doc.save(DOCX_PATH)
    print("DOCX successfully generated.")

def generate_pdf():
    print(f"Generating PDF at {PDF_PATH}...")
    doc = SimpleDocTemplate(
        PDF_PATH,
        pagesize=letter,
        rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=24,
        leading=28,
        textColor=colors.HexColor('#2b5590'),
        alignment=1, # Center
        spaceAfter=15
    )
    
    subtitle_style = ParagraphStyle(
        'DocSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=13,
        leading=16,
        textColor=colors.HexColor('#787878'),
        alignment=1,
        spaceAfter=40
    )
    
    meta_style = ParagraphStyle(
        'DocMeta',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#333333'),
        alignment=1,
        spaceAfter=200
    )
    
    h1_style = ParagraphStyle(
        'H1Style',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=15,
        leading=18,
        textColor=colors.HexColor('#2b5590'),
        spaceBefore=15,
        spaceAfter=8,
        keepWithNext=True
    )
    
    h2_style = ParagraphStyle(
        'H2Style',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        textColor=colors.HexColor('#333333'),
        spaceBefore=10,
        spaceAfter=4,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=13.5,
        textColor=colors.HexColor('#333333'),
        spaceAfter=8
    )
    
    code_style = ParagraphStyle(
        'CodeStyle',
        parent=styles['Normal'],
        fontName='Courier',
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor('#1e1e1e'),
        spaceAfter=10
    )
    
    curl_style = ParagraphStyle(
        'CurlStyle',
        parent=styles['Normal'],
        fontName='Courier-Bold',
        fontSize=8,
        leading=10.5,
        textColor=colors.HexColor('#1f3a60'),
        spaceAfter=10
    )
    
    story = []
    
    # Title Page
    story.append(Spacer(1, 100))
    story.append(Paragraph("Setu Reconciliation API Documentation", title_style))
    story.append(Paragraph("Developer API Reference & Integration Guide", subtitle_style))
    story.append(Spacer(1, 120))
    story.append(Paragraph("Backend Service: FastAPI<br/>Database: PostgreSQL<br/>Date: July 2026<br/>Version: 1.0.0", meta_style))
    story.append(PageBreak())
    
    # Introduction
    story.append(Paragraph("Overview", h1_style))
    story.append(Paragraph(
        "This service provides payment lifecycle event ingestion and reconciliation capabilities. "
        "It supports event logging with strict idempotency check, transaction state maintenance, "
        "and discrepancy checks for reporting. This document describes the schema structure, request "
        "payloads, query parameters, responses, and curl commands for all active API endpoints.",
        body_style
    ))
    story.append(Paragraph("Base URL", h2_style))
    story.append(Paragraph("Local Server: http://localhost:8000", body_style))
    story.append(PageBreak())
    
    # API endpoints
    for ep in endpoints:
        # Title of API
        story.append(Paragraph(ep["title"], h1_style))
        
        # Path & Method
        color_method = "#c83232" if ep["method"] == "POST" else "#329632"
        p_str = f"<b><font color='{color_method}'>[{ep['method']}]</font></b> <font size='11.5'><b>{ep['path']}</b></font>"
        story.append(Paragraph(p_str, body_style))
        story.append(Paragraph(ep["desc"], body_style))
        
        # Parameters Table
        if ep["params"]:
            story.append(Paragraph("Query Parameters", h2_style))
            data = [["Parameter", "Type", "Req", "Description"]]
            for name, ptype, req, pdesc in ep["params"]:
                data.append([name, ptype, req, pdesc])
                
            t = Table(data, colWidths=[100, 60, 40, 310])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#2b5590')),
                ('TEXTCOLOR', (0,0), (-1,0), colors.white),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0,0), (-1,0), 6),
                ('TOPPADDING', (0,0), (-1,0), 6),
                ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#dddddd')),
                ('FONTNAME', (0,1), (-1,-1), 'Helvetica'),
                ('FONTSIZE', (0,0), (-1,-1), 9),
                ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f9f9f9')]),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('BOTTOMPADDING', (0,1), (-1,-1), 4),
                ('TOPPADDING', (0,1), (-1,-1), 4),
            ]))
            story.append(t)
            story.append(Spacer(1, 10))
            
        # Request Body Block
        if ep["req_body"]:
            story.append(Paragraph("Request Payload (JSON)", h2_style))
            json_str = json.dumps(ep["req_body"], indent=2).replace(" ", "&nbsp;").replace("\n", "<br/>")
            t_box = Table([[Paragraph(json_str, code_style)]], colWidths=[510])
            t_box.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f5f5f5')),
                ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#cccccc')),
                ('TOPPADDING', (0,0), (-1,-1), 8),
                ('BOTTOMPADDING', (0,0), (-1,-1), 8),
                ('LEFTPADDING', (0,0), (-1,-1), 10),
                ('RIGHTPADDING', (0,0), (-1,-1), 10),
            ]))
            story.append(t_box)
            story.append(Spacer(1, 10))
            
        # Response Body Block
        if ep["resp_body"]:
            story.append(Paragraph("Response Payload (JSON)", h2_style))
            json_str = json.dumps(ep["resp_body"], indent=2).replace(" ", "&nbsp;").replace("\n", "<br/>")
            t_box = Table([[Paragraph(json_str, code_style)]], colWidths=[510])
            t_box.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f5f5f5')),
                ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#cccccc')),
                ('TOPPADDING', (0,0), (-1,-1), 8),
                ('BOTTOMPADDING', (0,0), (-1,-1), 8),
                ('LEFTPADDING', (0,0), (-1,-1), 10),
                ('RIGHTPADDING', (0,0), (-1,-1), 10),
            ]))
            story.append(t_box)
            story.append(Spacer(1, 10))
            
        # Curl Block
        story.append(Paragraph("Sample curl Command", h2_style))
        curl_str = ep["curl"].replace(" ", "&nbsp;").replace("\n", "<br/>")
        t_box_curl = Table([[Paragraph(curl_str, curl_style)]], colWidths=[510])
        t_box_curl.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#e8effc')),
            ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#b3cbf2')),
            ('TOPPADDING', (0,0), (-1,-1), 8),
            ('BOTTOMPADDING', (0,0), (-1,-1), 8),
            ('LEFTPADDING', (0,0), (-1,-1), 10),
            ('RIGHTPADDING', (0,0), (-1,-1), 10),
        ]))
        story.append(t_box_curl)
        
        story.append(PageBreak())
        
    doc.build(story)
    print("PDF successfully generated.")

if __name__ == "__main__":
    generate_docx()
    generate_pdf()
    print("Documentation generation completed successfully!")
